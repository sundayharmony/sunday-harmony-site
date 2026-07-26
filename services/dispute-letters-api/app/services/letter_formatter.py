from __future__ import annotations

import html
import re
from datetime import date


SECTION_HEADINGS = frozenset(
    {
        "Consumer Identification",
        "Disputed Tradelines",
        "Statutory Reinvestigation Requirements",
        "Requested Outcome",
        "CONSUMER INFORMATION",
        "DISPUTED ITEMS",
        "STATUTORY REINVESTIGATION REQUIREMENTS",
        "REQUESTED OUTCOME",
    }
)

_BULLET_RE = re.compile(r"^(?:●|•|▪|◦|[-*+])\s+")
_FIELD_LABEL_RE = re.compile(
    r"^(Full Name|Date of Birth|Current Address|Additional Addresses on File|"
    r"Account Number|Reported Status|Reported Balance|Basis of Dispute)\s*:",
    re.IGNORECASE,
)
_DATE_LINE_RE = re.compile(
    r"^(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},\s+\d{4}\s*$"
)


def normalize_letter_source(text: str) -> str:
    """Clean agent output while preserving controlled **bold** markup."""
    text = text.strip()
    text = re.sub(r"```[\w]*\n?", "", text)
    text = text.replace("```", "")
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^[-*_]{3,}\s*$", "", text, flags=re.MULTILINE)
    # Normalize bullet markers to a filled circle for consistent print layout
    text = re.sub(r"^(\s*)(?:[•▪◦]|\-|\*|\+)\s+", r"\1● ", text, flags=re.MULTILINE)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def strip_markdown(text: str) -> str:
    """Plain text for .txt download — no bold markers."""
    text = normalize_letter_source(text)
    text = re.sub(r"\*\*\*(.+?)\*\*\*", r"\1", text)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"_(.+?)_", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()


def finalize_letter(text: str) -> str:
    cleaned = strip_markdown(text)
    cleaned = re.sub(
        r"\n+This letter was prepared with automated assistance.*$",
        "",
        cleaned,
        flags=re.IGNORECASE | re.DOTALL,
    )
    return cleaned.strip()


def _indent_level(line: str) -> int:
    if not line.strip():
        return 0
    leading = len(line) - len(line.lstrip(" "))
    if leading >= 8:
        return 2
    if leading >= 4:
        return 1
    return 0


def _line_display(line: str) -> str:
    return line.strip()


def _is_bullet(line: str) -> bool:
    return bool(_BULLET_RE.match(line.strip()))


def _bullet_text(line: str) -> str:
    return _BULLET_RE.sub("", line.strip()).strip()


def _is_section_heading(line: str) -> bool:
    stripped = line.strip()
    if not stripped or stripped.endswith(":") or "**" in stripped:
        return False
    if stripped in SECTION_HEADINGS:
        return True
    if any(c.isdigit() for c in stripped):
        return False
    # Do not treat ALL-CAPS consumer names as headings — only known section labels above.
    if stripped.isupper():
        return False
    words = stripped.split()
    if len(words) < 2 or len(stripped) > 55:
        return False
    small = {"of", "and", "the", "or", "to", "a", "an"}
    return all(w[:1].isupper() for w in words if w.lower() not in small)


def _is_name_line(line: str, *, index: int) -> bool:
    """ALL-CAPS consumer name typically appears near the top after the date."""
    stripped = line.strip()
    if index > 6 or not stripped or len(stripped) > 60:
        return False
    if any(c.isdigit() for c in stripped):
        return False
    letters = [c for c in stripped if c.isalpha()]
    if len(letters) < 4:
        return False
    return stripped == stripped.upper() and " " in stripped


def _is_field_label_line(line: str) -> bool:
    return bool(_FIELD_LABEL_RE.match(line.strip()))


def _is_subheading(line: str) -> bool:
    return line.strip().lower().rstrip(":") == "basis of dispute"


def _extract_letter_date(body: str) -> str:
    for line in body.splitlines()[:8]:
        if _DATE_LINE_RE.match(line.strip()):
            return line.strip()
    return date.today().strftime("%B %d, %Y")


def _inline_html(text: str) -> str:
    """Render **bold** spans inside escaped text."""
    parts = re.split(r"(\*\*.+?\*\*)", text)
    out: list[str] = []
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            inner = html.escape(part[2:-2])
            out.append(f"<strong>{inner}</strong>")
        else:
            out.append(html.escape(part))
    return "".join(out)


def letter_to_html(text: str) -> str:
    body = normalize_letter_source(text)
    lines = body.splitlines()
    blocks: list[str] = []
    for i, line in enumerate(lines):
        if not line.strip():
            blocks.append('<p class="letter-spacer">&nbsp;</p>')
            continue
        display = _line_display(line)
        if _is_bullet(display):
            inner = _inline_html(_bullet_text(display))
            blocks.append(f'<p class="letter-line letter-bullet">● {inner}</p>')
            continue
        indent = _indent_level(line)
        inner = _inline_html(display)
        classes = ["letter-line"]
        if indent:
            classes.append(f"letter-indent-{indent}")
        if _is_section_heading(display):
            classes.append("letter-heading")
        elif _is_name_line(display, index=i):
            classes.append("letter-name")
        elif _is_field_label_line(display) or _is_subheading(display):
            classes.append("letter-field")
        blocks.append(f'<p class="{" ".join(classes)}">{inner}</p>')
    return "\n".join(blocks)


def _add_runs(para, line: str, *, bold_all: bool = False) -> None:
    parts = re.split(r"(\*\*.+?\*\*)", line)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            run = para.add_run(part)
            if bold_all:
                run.bold = True


def _add_rich_paragraph(
    document,
    line: str,
    *,
    indent_level: int = 0,
    heading: bool = False,
    name_line: bool = False,
    field_line: bool = False,
    tight: bool = False,
):
    from docx.shared import Inches, Pt

    para = document.add_paragraph()
    fmt = para.paragraph_format
    fmt.space_before = Pt(12) if heading else Pt(0)
    if heading:
        fmt.space_after = Pt(8)
    elif name_line or field_line or tight:
        fmt.space_after = Pt(2)
    else:
        fmt.space_after = Pt(8)

    if indent_level:
        fmt.left_indent = Inches(0.35 * indent_level)

    _add_runs(para, line, bold_all=heading or name_line)
    return para


def _add_bullet_paragraph(document, text: str):
    from docx.shared import Inches, Pt

    para = document.add_paragraph()
    fmt = para.paragraph_format
    fmt.left_indent = Inches(0.25)
    fmt.first_line_indent = Inches(-0.2)
    fmt.space_after = Pt(4)
    fmt.space_before = Pt(0)
    _add_runs(para, f"● {text}")
    return para


def _set_run_font(run, *, size_pt: float = 11) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")


def _add_page_field(paragraph, instr: str):
    """Insert a Word field such as PAGE or NUMPAGES."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    _set_run_font(run, size_pt=11)
    r = run._r

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f" {instr} "

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    # Placeholder shown until Word updates fields
    placeholder = OxmlElement("w:t")
    placeholder.text = "1" if instr == "PAGE" else "1"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r.append(fld_begin)
    r.append(instr_text)
    r.append(fld_separate)
    r.append(placeholder)
    r.append(fld_end)
    return run


def letter_to_docx(text: str):
    from docx import Document
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    body = normalize_letter_source(text)
    letter_date = _extract_letter_date(body)
    document = Document()

    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.clear()
        hp.paragraph_format.space_after = Pt(6)
        tab_stops = hp.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)

        run_date = hp.add_run(letter_date)
        _set_run_font(run_date, size_pt=11)
        hp.add_run("\t")
        run_page = hp.add_run("Page ")
        _set_run_font(run_page, size_pt=11)
        _add_page_field(hp, "PAGE")
        slash = hp.add_run("/")
        _set_run_font(slash, size_pt=11)
        _add_page_field(hp, "NUMPAGES")

    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    lines = body.splitlines()
    start = 0
    if lines and _DATE_LINE_RE.match(lines[0].strip()):
        start = 1
        if start < len(lines) and not lines[start].strip():
            start += 1

    prev_blank = True
    for i, line in enumerate(lines[start:], start=start):
        if not line.strip():
            if not prev_blank:
                document.add_paragraph("")
            prev_blank = True
            continue
        prev_blank = False
        display = _line_display(line)

        if _is_bullet(display):
            _add_bullet_paragraph(document, _bullet_text(display))
            continue

        heading = _is_section_heading(display)
        name_line = _is_name_line(display, index=i)
        field_line = _is_field_label_line(display) or _is_subheading(display)
        tight = (
            (not heading and not name_line and len(display) < 70 and ":" not in display)
            or field_line
        )

        _add_rich_paragraph(
            document,
            display,
            indent_level=_indent_level(line),
            heading=heading,
            name_line=name_line,
            field_line=field_line,
            tight=tight,
        )

    return document
