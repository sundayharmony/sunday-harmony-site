from __future__ import annotations

from pathlib import Path

from app.models import ExtractedDocument


def _quality_from_chars(count: int) -> str:
    if count >= 5000:
        return "high"
    if count >= 1000:
        return "medium"
    return "low"


def extract_docx(path: Path) -> ExtractedDocument:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    return ExtractedDocument(
        file_type="docx",
        raw_path=str(path),
        text=text,
        html=None,
        ocr_used=False,
        extraction_quality=_quality_from_chars(len(text)),
    )
